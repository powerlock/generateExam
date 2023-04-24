import random
from docx import Document
import re
# Open the Word document
doc = Document('chapter9.docx')

# Define the bank of questions with their levels of difficulty
question_bank = {
    "G1": [],
    "G2": [],
    "G3": [],
    "G4": []
}

# Initialize variables to keep track of the current question and level
current_question = None
current_level = None
opt =[]

# Iterate through each paragraph in the Word document
for paragraph in doc.paragraphs:
    # Extract the text from the paragraph
    text = paragraph.text
    
    # Check if the paragraph starts with a number followed by a dot, indicating a new question
    if text and text[0].isdigit() and ')' in text:
        # Extract the question from the text
        question = re.sub(r'^\d+\)', '', text).strip()#text.split(')')[1].strip()
        current_question = question
    # Check if the paragraph starts with a A, B, C or D followed by a half parenthesis, indicating a new option
    elif text and text[0].isalpha() and ')' in text:
    #elif any(option in text for option in ['A)', 'B)', 'C)', 'D)']):
        
        # find the options in the text
        options = re.sub(r'^\d+\)', '', text).strip()#text.split(')')[1].strip()
        if current_question is not None:
            # add the options to the list
            opt.append(options)
            print(opt)
        
    #Check if the paragraph contains the answer
    elif 'Answer:' in text:
        match = re.search(r'Answer:\s*([A-D])', text)
        if match:
            answer = match.group(1)
        else:
            answer = 'N/A'


    # Check if the paragraph contains the level keyword in the description at the bottom
    elif 'Global Obj:' in text:
        
        # Extract the level keyword from the text
        level = text.split('G')[-1].strip().split()[0]
        current_level = 'G'+str(level)
        # Add the current question to the question bank with the current level
        if current_question is not None and current_level is not None:
            
                    # Add the current question to the question bank with the options, answer and current level
            question_bank[current_level].append({
            "question": current_question,
            "options": opt,
            
        })
            current_question = None
            current_level = None
            opt=[]
            answer = None
        

# Input the number of questions for each level of difficulty
# num_g1 = int(input("Enter the number of G1 questions: "))
# num_g2 = int(input("Enter the number of G2 questions: "))
# num_g3 = int(input("Enter the number of G3 questions: "))
num_g4 = int(input("Enter the number of G4 questions: "))

# Initialize the list to store the selected questions
selected_questions = []

# Randomly select questions from the question bank based on the specified number of questions for each level of difficulty
for level, questions in question_bank.items():
    # if level == "G1":
    #     selected_questions.extend(random.sample(questions, num_g1))
    # elif level == "G2":
    #     selected_questions.extend(random.sample(questions, num_g2))
    # elif level == "G3":
    #     selected_questions.extend(random.sample(questions, num_g3))
    if level == "G4":
        selected_questions.extend(random.sample(questions, num_g4))

# Create a new Word document to store the selected questions
output_doc = Document()
# Write the selected questions to the new Word document
for i, Question in enumerate(selected_questions, 1):
    # Add a new paragraph for each question
    question_text = Question['question']
    options = Question['options']
    output_doc.add_paragraph(f"Question {i}: {question_text}")
    for option in options:
        output_doc.add_paragraph(f"{option}")

# Save the new Word document
output_doc.save('generated_exam.docx')

print("Generated exam has been saved to 'generated_exam.docx'.")
